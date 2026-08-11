import shutil
import tempfile
from datetime import timedelta
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import override_settings
from django.utils import timezone
from docx import Document
from pypdf import PdfWriter
from rest_framework.test import APITestCase

from .flight_permits.services.lifecycle import (
    create_flight_permit,
    update_flight_permit,
)
from .models import FlightPermit


def valid_pdf_bytes() -> bytes:
    output = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.write(output)
    return output.getvalue()


class FlightPermitApiTests(APITestCase):
    def setUp(self):
        self.media_root = tempfile.mkdtemp(prefix="uav-flight-permits-")
        self.media_override = override_settings(MEDIA_ROOT=self.media_root)
        self.media_override.enable()
        self.user = get_user_model().objects.create_user(
            username="permit-user",
            password="test-password",
            is_active=True,
        )
        self.client.force_authenticate(self.user)

    def tearDown(self):
        self.media_override.disable()
        shutil.rmtree(self.media_root, ignore_errors=True)
        super().tearDown()

    def permit_payload(self, **overrides):
        today = timezone.localdate()
        payload = {
            "aircraft_number": "tc-uav-104",
            "permit_number": "shgm-ui-2026-0042",
            "permit_type": FlightPermit.TYPE_TEST,
            "issuing_authority": "SHGM",
            "flight_region": "Ankara Test Sahası",
            "valid_from": today.isoformat(),
            "valid_until": (today + timedelta(days=60)).isoformat(),
            "status": FlightPermit.STATUS_APPROVED,
            "notes": "Gündüz VFR operasyonları",
        }
        payload.update(overrides)
        return payload

    def test_create_list_update_and_open_document(self):
        upload = SimpleUploadedFile(
            "ucus-izni.pdf",
            valid_pdf_bytes(),
            content_type="application/pdf",
        )
        response = self.client.post(
            "/api/flight-permits/",
            {**self.permit_payload(), "document": upload},
            format="multipart",
        )

        self.assertEqual(response.status_code, 201)
        self.assertEqual(response.data["aircraft_number"], "TC-UAV-104")
        self.assertEqual(response.data["permit_number"], "SHGM-UI-2026-0042")
        self.assertEqual(response.data["validity_status"], "active")
        self.assertEqual(response.data["document_name"], "ucus-izni.pdf")
        self.assertTrue(response.data["document_url"].endswith("/api/flight-permits/1/document/"))
        self.assertTrue(
            response.data["generated_document_url"].endswith(
                "/api/flight-permits/1/generated-document/"
            )
        )

        permit = FlightPermit.objects.get()
        stored_document = Path(permit.document.path)
        self.assertTrue(stored_document.exists())
        self.assertEqual(permit.created_by, self.user)

        document_response = self.client.get(f"/api/flight-permits/{permit.id}/document/")
        self.assertEqual(document_response.status_code, 200)
        self.assertEqual(document_response["Content-Type"], "application/pdf")

        with self.captureOnCommitCallbacks(execute=True):
            update_response = self.client.patch(
                f"/api/flight-permits/{permit.id}/",
                {"flight_region": "Konya Test Sahası", "remove_document": "true"},
                format="multipart",
            )
        self.assertEqual(update_response.status_code, 200)
        self.assertEqual(update_response.data["flight_region"], "Konya Test Sahası")
        self.assertEqual(update_response.data["document_url"], "")
        self.assertFalse(stored_document.exists())

        list_response = self.client.get("/api/flight-permits/")
        self.assertEqual(list_response.status_code, 200)
        self.assertEqual(len(list_response.data), 1)

    def test_generates_downloadable_word_permit_from_record(self):
        create_response = self.client.post(
            "/api/flight-permits/",
            self.permit_payload(),
            format="multipart",
        )
        permit_id = create_response.data["id"]

        response = self.client.get(f"/api/flight-permits/{permit_id}/generated-document/")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertIn("attachment", response["Content-Disposition"])
        self.assertIn(
            "Ucus_Izni_TC-UAV-104_SHGM-UI-2026-0042.docx", response["Content-Disposition"]
        )

        generated = Document(BytesIO(b"".join(response.streaming_content)))
        text_parts = [paragraph.text for paragraph in generated.paragraphs]
        for table in generated.tables:
            text_parts.extend(cell.text for row in table.rows for cell in row.cells)
        generated_text = "\n".join(text_parts)
        self.assertIn("UÇUŞ İZNİ", generated_text)
        self.assertIn("TC-UAV-104", generated_text)
        self.assertIn("SHGM-UI-2026-0042", generated_text)
        self.assertIn("Ankara Test Sahası", generated_text)
        self.assertNotIn("{{", generated_text)

    def test_rejects_invalid_date_range_and_document_type(self):
        today = timezone.localdate()
        invalid_dates = self.client.post(
            "/api/flight-permits/",
            self.permit_payload(
                valid_from=(today + timedelta(days=3)).isoformat(),
                valid_until=today.isoformat(),
            ),
            format="multipart",
        )
        self.assertEqual(invalid_dates.status_code, 400)
        self.assertIn("valid_until", invalid_dates.data)

        invalid_file = SimpleUploadedFile("izin.exe", b"not allowed")
        invalid_document = self.client.post(
            "/api/flight-permits/",
            {
                **self.permit_payload(permit_number="SHGM-UI-2026-0043"),
                "document": invalid_file,
            },
            format="multipart",
        )
        self.assertEqual(invalid_document.status_code, 400)
        self.assertIn("document", invalid_document.data)

        for index, filename in enumerate(("izin.doc", "liste.xls"), start=4):
            with self.subTest(filename=filename):
                legacy_document = self.client.post(
                    "/api/flight-permits/",
                    {
                        **self.permit_payload(permit_number=f"SHGM-UI-2026-004{index}"),
                        "document": SimpleUploadedFile(filename, b"legacy binary"),
                    },
                    format="multipart",
                )
                self.assertEqual(legacy_document.status_code, 400)
                self.assertIn("document", legacy_document.data)

    def test_rejects_content_that_does_not_match_the_permitted_extension(self):
        response = self.client.post(
            "/api/flight-permits/",
            {
                **self.permit_payload(),
                "document": SimpleUploadedFile(
                    "izin.pdf",
                    b"<script>alert('xss')</script>",
                    content_type="text/html",
                ),
            },
            format="multipart",
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("document", response.data)
        self.assertFalse(FlightPermit.objects.exists())

    def test_server_mime_policy_ignores_spoofed_client_content_type(self):
        response = self.client.post(
            "/api/flight-permits/",
            {
                **self.permit_payload(),
                "document": SimpleUploadedFile(
                    "izin.pdf",
                    valid_pdf_bytes(),
                    content_type="text/html",
                ),
            },
            format="multipart",
        )

        self.assertEqual(response.status_code, 201)
        self.assertEqual(response.data["document_content_type"], "application/pdf")
        document_response = self.client.get(f"/api/flight-permits/{response.data['id']}/document/")
        self.assertEqual(document_response.status_code, 200)
        self.assertEqual(document_response["Content-Type"], "application/pdf")
        self.assertEqual(document_response["X-Content-Type-Options"], "nosniff")

    def test_deleting_permit_deletes_stored_document(self):
        upload = SimpleUploadedFile("izin.pdf", valid_pdf_bytes())
        create_response = self.client.post(
            "/api/flight-permits/",
            {**self.permit_payload(), "document": upload},
            format="multipart",
        )
        permit = FlightPermit.objects.get(pk=create_response.data["id"])
        stored_document = Path(permit.document.path)

        with self.captureOnCommitCallbacks(execute=True):
            delete_response = self.client.delete(f"/api/flight-permits/{permit.id}/")

        self.assertEqual(delete_response.status_code, 204)
        self.assertFalse(FlightPermit.objects.filter(pk=permit.id).exists())
        self.assertFalse(stored_document.exists())

    def test_create_failure_rolls_back_record_and_compensates_stored_document(self):
        original_save = FlightPermit.save

        def save_then_fail(instance, *args, **kwargs):
            original_save(instance, *args, **kwargs)
            raise RuntimeError("database commit failed")

        today = timezone.localdate()
        with (
            patch.object(FlightPermit, "save", new=save_then_fail),
            self.assertRaisesMessage(RuntimeError, "database commit failed"),
        ):
            create_flight_permit(
                validated_data={
                    "aircraft_number": "TC-UAV-FAIL",
                    "permit_number": "SHGM-FAIL-1",
                    "permit_type": FlightPermit.TYPE_TEST,
                    "issuing_authority": "SHGM",
                    "valid_from": today,
                    "valid_until": today + timedelta(days=60),
                    "status": FlightPermit.STATUS_APPROVED,
                    "document": SimpleUploadedFile(
                        "rollback.pdf",
                        b"%PDF-1.4 rollback",
                        content_type="application/pdf",
                    ),
                },
                actor=self.user,
            )

        self.assertFalse(FlightPermit.objects.exists())
        self.assertEqual(list(Path(self.media_root).rglob("rollback.pdf")), [])

    def test_update_failure_keeps_old_file_and_compensates_replacement(self):
        create_response = self.client.post(
            "/api/flight-permits/",
            {
                **self.permit_payload(),
                "document": SimpleUploadedFile(
                    "original.pdf",
                    valid_pdf_bytes(),
                    content_type="application/pdf",
                ),
            },
            format="multipart",
        )
        permit = FlightPermit.objects.get(pk=create_response.data["id"])
        original_name = permit.document.name
        original_path = Path(permit.document.path)
        original_save = FlightPermit.save

        def save_then_fail(instance, *args, **kwargs):
            original_save(instance, *args, **kwargs)
            raise RuntimeError("database update failed")

        with (
            patch.object(FlightPermit, "save", new=save_then_fail),
            self.assertRaisesMessage(RuntimeError, "database update failed"),
        ):
            update_flight_permit(
                permit=permit,
                validated_data={
                    "document": SimpleUploadedFile(
                        "replacement.pdf",
                        b"%PDF-1.4 replacement",
                        content_type="application/pdf",
                    )
                },
                actor=self.user,
            )

        persisted = FlightPermit.objects.get(pk=permit.pk)
        self.assertEqual(persisted.document.name, original_name)
        self.assertTrue(original_path.exists())
        self.assertEqual(list(Path(self.media_root).rglob("replacement.pdf")), [])
