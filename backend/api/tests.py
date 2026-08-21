from datetime import timedelta
from io import BytesIO
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.contrib.auth.models import Group
from django.core import mail
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import Client, TestCase, override_settings
from django.urls import reverse
from django.utils import timezone

from .edk.minutes_parser import EDKMinutesParseError
from .edk.roles import EDK_ROLE_GROUPS
from .models import (
    CoverPage,
    EDKApplication,
    PanelResponsible,
    Person,
    PersonGroup,
    Project,
    ProjectPanel,
    TechnicalDocument,
    TechnicalDocumentNotification,
    TechnicalDocumentStatusHistory,
)


class AuthApiTests(TestCase):
    def setUp(self):
        self.client = Client(enforce_csrf_checks=True)

    def get_csrf_token(self):
        response = self.client.get(reverse("csrf-token"))
        self.assertEqual(response.status_code, 200)
        return response.json()["csrfToken"]

    def test_register_creates_pending_user_with_csrf(self):
        csrf_token = self.get_csrf_token()

        response = self.client.post(
            reverse("register"),
            data={
                "username": "operator",
                "email": "operator@example.com",
                "password": "StrongPass123!",
                "password_confirm": "StrongPass123!",
            },
            content_type="application/json",
            HTTP_X_CSRFTOKEN=csrf_token,
        )

        self.assertEqual(response.status_code, 201)
        self.assertFalse(response.json()["authenticated"])
        self.assertEqual(response.json()["user"]["username"], "operator")
        self.assertEqual(response.json()["user"]["email"], "operator@example.com")
        self.assertFalse(response.json()["user"]["is_active"])

        user = get_user_model().objects.get(username="operator")
        self.assertFalse(user.is_active)

    def test_login_requires_valid_credentials_and_csrf(self):
        user_model = get_user_model()
        user_model.objects.create_user(username="pilot", password="StrongPass123!")
        csrf_token = self.get_csrf_token()

        response = self.client.post(
            reverse("login"),
            data={"username": "pilot", "password": "StrongPass123!"},
            content_type="application/json",
            HTTP_X_CSRFTOKEN=csrf_token,
        )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json()["authenticated"])
        self.assertEqual(response.json()["user"]["username"], "pilot")

    def test_pending_user_cannot_login_until_admin_approves(self):
        user_model = get_user_model()
        admin = user_model.objects.create_user(
            username="admin",
            email="admin@example.com",
            password="StrongPass123!",
            is_staff=True,
        )
        user = user_model.objects.create_user(
            username="candidate",
            email="candidate@example.com",
            password="StrongPass123!",
            is_active=False,
        )
        csrf_token = self.get_csrf_token()

        pending_login = self.client.post(
            reverse("login"),
            data={"username": "candidate", "password": "StrongPass123!"},
            content_type="application/json",
            HTTP_X_CSRFTOKEN=csrf_token,
        )

        self.assertEqual(pending_login.status_code, 400)

        self.client.force_login(admin)
        approval = self.client.patch(
            reverse("admin-user-status", kwargs={"user_id": user.id}),
            data={"is_active": True},
            content_type="application/json",
            HTTP_X_CSRFTOKEN=csrf_token,
        )

        self.assertEqual(approval.status_code, 200)
        self.assertTrue(approval.json()["is_active"])

        self.client.logout()
        approved_login = self.client.post(
            reverse("login"),
            data={"username": "candidate", "password": "StrongPass123!"},
            content_type="application/json",
            HTTP_X_CSRFTOKEN=csrf_token,
        )

        self.assertEqual(approved_login.status_code, 200)
        self.assertTrue(approved_login.json()["authenticated"])

    def test_documents_require_authenticated_session(self):
        response = self.client.get(reverse("document-list"))

        self.assertEqual(response.status_code, 403)

    def test_deactivated_session_cannot_access_documents(self):
        user_model = get_user_model()
        user = user_model.objects.create_user(
            username="operator",
            email="operator@example.com",
            password="StrongPass123!",
        )
        self.client.force_login(user)

        user.is_active = False
        user.save(update_fields=["is_active"])

        response = self.client.get(reverse("document-list"))

        self.assertEqual(response.status_code, 403)


class EDKMinutesParseApiTests(TestCase):
    def setUp(self):
        user = get_user_model().objects.create_user(
            username="edk-reader",
            password="StrongPass123!",
        )
        Group.objects.get_or_create(name=EDK_ROLE_GROUPS["applicant"])[0].user_set.add(user)
        self.application = EDKApplication.objects.create(
            applicant=user,
            meeting_title="Uçuş hazırlığı",
            project_name="UAV",
            requested_date=timezone.localdate(),
            location="Hangar",
            participants="Uçuş ekibi",
            purpose="Uçuş hazırlığını değerlendirmek",
            agenda="Hazırlık kontrolleri",
            status=EDKApplication.STATUS_APPROVED,
        )
        self.client.force_login(user)
        self.parse_url = reverse(
            "edk-minutes-parse",
            kwargs={"application_id": self.application.id},
        )

    @staticmethod
    def word_file():
        from docx import Document

        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Başlık"
        table.cell(0, 1).text = "Değer"
        table.cell(1, 0).text = "Özet"
        table.cell(1, 1).text = "Uçuş kontrolü"
        table.cell(0, 0).merge(table.cell(0, 1))
        content = BytesIO()
        document.save(content)
        return SimpleUploadedFile(
            "gorev.docx",
            content.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def test_parse_returns_zero_based_cell_coordinates(self):
        response = self.client.post(
            self.parse_url,
            data={"file": self.word_file()},
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["table_count"], 1)
        self.assertEqual(payload["cell_count"], 3)
        self.assertEqual(
            payload["cells"][2],
            {
                "index": 2,
                "table_index": 0,
                "row_index": 1,
                "column_index": 1,
                "text": "Uçuş kontrolü",
            },
        )
        self.assertEqual(payload["cells"][0]["text"], "Başlık\nDeğer")
        self.assertFalse(payload["jira_ready"])

    def test_parse_rejects_non_docx_file(self):
        response = self.client.post(
            self.parse_url,
            data={"file": SimpleUploadedFile("notlar.txt", b"test")},
        )

        self.assertEqual(response.status_code, 400)

    @override_settings(DOCUMENT_MAX_UPLOAD_SIZE=4)
    @patch("api.edk.views.parse_minutes_document")
    def test_parse_rejects_oversized_docx_before_temporary_write(self, parse_minutes_document):
        response = self.client.post(
            self.parse_url,
            data={"file": self.word_file()},
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("file", response.json())
        parse_minutes_document.assert_not_called()

    @patch("api.edk.views.parse_minutes_document")
    def test_parse_rejects_invalid_ooxml_before_temporary_write(self, parse_minutes_document):
        response = self.client.post(
            self.parse_url,
            data={
                "file": SimpleUploadedFile(
                    "invalid.docx",
                    b"not an OOXML archive",
                    content_type=(
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ),
                )
            },
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("file", response.json())
        parse_minutes_document.assert_not_called()

    @patch("api.edk.views.parse_minutes_document")
    def test_parse_error_does_not_echo_temporary_path(self, parse_minutes_document):
        parse_minutes_document.side_effect = EDKMinutesParseError(
            "invalid /private/tmp/upload-secret.docx pilot@example.com"
        )

        response = self.client.post(
            self.parse_url,
            data={"file": self.word_file()},
        )

        self.assertEqual(response.status_code, 422)
        self.assertEqual(response.json()["detail"], "Word dosyası işlenemedi.")
        self.assertNotIn("/private/tmp", response.content.decode())

    def test_parse_keeps_all_distinct_cells(self):
        from docx import Document

        document = Document()
        table = document.add_table(rows=10, cols=5)
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                cell.text = f"{row_index}:{column_index}"

        content = BytesIO()
        document.save(content)
        upload = SimpleUploadedFile(
            "buyuk-tablo.docx",
            content.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        response = self.client.post(
            self.parse_url,
            data={"file": upload},
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["cell_count"], 50)
        self.assertEqual(payload["cells"][-1]["text"], "9:4")

    def test_extracts_mapped_fields_and_dynamic_action_item_range(self):
        from docx import Document

        document = Document()
        table = document.add_table(rows=15, cols=4)
        table.cell(1, 2).text = "UAV"
        table.cell(2, 2).text = "Uçuş hazırlığı"
        table.cell(3, 2).text = "MOM-42"
        table.cell(4, 2).text = "B"
        table.cell(8, 0).merge(table.cell(8, 3)).text = "Action Item List"
        headers = ["No", "Action Item", "Responsible", "Due Date"]
        for column_index, header in enumerate(headers):
            table.cell(9, column_index).text = header
        first_item = ["1", "Motor kontrolünü tamamla", "Ada", "2026-07-10"]
        second_item = ["2", "Telemetri raporunu paylaş", "Deniz", "2026-07-12"]
        for column_index, value in enumerate(first_item):
            table.cell(10, column_index).text = value
        for column_index, value in enumerate(second_item):
            table.cell(11, column_index).text = value
        table.cell(12, 0).merge(table.cell(12, 3)).text = "Attachments / Ekler"
        table.cell(13, 0).text = "action-items.pdf"

        content = BytesIO()
        document.save(content)
        upload = SimpleUploadedFile(
            "mom.docx",
            content.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        response = self.client.post(
            self.parse_url,
            data={"file": upload},
        )

        self.assertEqual(response.status_code, 200)
        extracted = response.json()["extracted_data"]
        self.assertEqual(extracted["project"], "UAV")
        self.assertEqual(extracted["subject"], "Uçuş hazırlığı")
        self.assertEqual(extracted["mom_no"], "MOM-42")
        self.assertEqual(extracted["revision"], "B")
        self.assertEqual(
            extracted["action_items"],
            [
                {
                    "no": "1",
                    "action_item": "Motor kontrolünü tamamla",
                    "responsible": "Ada",
                    "due_date": "2026-07-10",
                },
                {
                    "no": "2",
                    "action_item": "Telemetri raporunu paylaş",
                    "responsible": "Deniz",
                    "due_date": "2026-07-12",
                },
            ],
        )
        self.assertTrue(extracted["action_item_list_found"])
        self.assertTrue(extracted["attachments_found"])


class OrganizationApiTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.user = user_model.objects.create_user(username="viewer", password="StrongPass123!")
        self.admin = user_model.objects.create_user(
            username="organization-admin",
            password="StrongPass123!",
            is_staff=True,
        )
        self.project = Project.objects.create(name="Uçuş Sistemleri", code="UAV")
        self.panel = ProjectPanel.objects.create(project=self.project, name="Aviyonik")
        PanelResponsible.objects.create(
            panel=self.panel,
            name="Ada Yılmaz",
            title="Panel Lideri",
            email="ada@example.com",
        )

    def test_authenticated_user_sees_nested_organization(self):
        self.client.force_login(self.user)
        response = self.client.get(reverse("project-list"))

        self.assertEqual(response.status_code, 200)
        project = response.json()[0]
        self.assertEqual(project["code"], "UAV")
        self.assertEqual(project["panels"][0]["name"], "Aviyonik")
        self.assertEqual(project["panels"][0]["responsibles"][0]["name"], "Ada Yılmaz")

    def test_regular_user_cannot_change_organization(self):
        self.client.force_login(self.user)
        response = self.client.post(
            reverse("project-list"),
            data={"name": "Yeni Proje", "code": "NEW"},
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 403)
        self.assertFalse(Project.objects.filter(code="NEW").exists())

    def test_admin_can_manage_all_organization_levels(self):
        self.client.force_login(self.admin)
        project_response = self.client.post(
            reverse("project-list"),
            data={"name": "Yer İstasyonu", "code": "GCS"},
            content_type="application/json",
        )
        self.assertEqual(project_response.status_code, 201)

        panel_response = self.client.post(
            reverse("project-panel-list", kwargs={"project_id": project_response.json()["id"]}),
            data={"name": "Haberleşme"},
            content_type="application/json",
        )
        self.assertEqual(panel_response.status_code, 201)

        responsible_response = self.client.post(
            reverse("panel-responsible-list", kwargs={"panel_id": panel_response.json()["id"]}),
            data={"name": "Deniz Kaya", "email": "deniz@example.com"},
            content_type="application/json",
        )
        self.assertEqual(responsible_response.status_code, 201)
        self.assertEqual(responsible_response.json()["order"], 0)

        second_responsible_response = self.client.post(
            reverse("panel-responsible-list", kwargs={"panel_id": panel_response.json()["id"]}),
            data={"name": "Ece Arslan", "order": 99},
            content_type="application/json",
        )
        self.assertEqual(second_responsible_response.status_code, 201)
        self.assertEqual(second_responsible_response.json()["order"], 1)

        delete_response = self.client.delete(
            reverse("project-detail", kwargs={"project_id": project_response.json()["id"]})
        )
        self.assertEqual(delete_response.status_code, 204)
        self.assertFalse(ProjectPanel.objects.filter(id=panel_response.json()["id"]).exists())

    def test_authenticated_user_sees_person_groups_but_cannot_change_them(self):
        group = PersonGroup.objects.create(name="Uçuş Ekibi", description="Test uçuşları")
        person = Person.objects.create(name="Selin Ak", email="selin@example.com")
        group.people.add(person)
        self.client.force_login(self.user)

        response = self.client.get(reverse("person-group-list"))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()[0]["name"], "Uçuş Ekibi")
        self.assertEqual(response.json()[0]["people"][0]["name"], "Selin Ak")

        create_response = self.client.post(
            reverse("person-group-list"),
            data={"name": "Yetkisiz Grup"},
            content_type="application/json",
        )
        self.assertEqual(create_response.status_code, 403)

    def test_admin_creates_group_and_person(self):
        self.client.force_login(self.admin)
        group_response = self.client.post(
            reverse("person-group-list"),
            data={"name": "Kalite Ekibi", "description": "Kalite güvence"},
            content_type="application/json",
        )
        self.assertEqual(group_response.status_code, 201)

        person_response = self.client.post(
            reverse("group-person-list", kwargs={"group_id": group_response.json()["id"]}),
            data={
                "name": "Mert Can",
                "title": "Kalite Mühendisi",
                "email": "mert@example.com",
            },
            content_type="application/json",
        )
        self.assertEqual(person_response.status_code, 201)
        self.assertEqual(person_response.json()["groups"], [group_response.json()["id"]])
        self.assertTrue(
            PersonGroup.objects.get(pk=group_response.json()["id"])
            .people.filter(pk=person_response.json()["id"])
            .exists()
        )


@override_settings(EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend")
class TechnicalDocumentApiTests(TestCase):
    def setUp(self):
        user_model = get_user_model()
        self.user = user_model.objects.create_user(username="engineer", password="StrongPass123!")
        self.admin = user_model.objects.create_user(
            username="document-admin",
            password="StrongPass123!",
            is_staff=True,
        )
        self.project = Project.objects.create(name="TULPAR", code="TPL")
        self.panel = ProjectPanel.objects.create(project=self.project, name="Aviyonik")
        self.other_project = Project.objects.create(name="GÖKBEY", code="GKB")
        self.other_panel = ProjectPanel.objects.create(project=self.other_project, name="Yapısal")
        PanelResponsible.objects.create(
            panel=self.panel,
            name="Ada Yılmaz",
            title="Panel Lideri",
            email="ada@example.com",
        )

    def create_document(self, **overrides):
        values = {
            "project": self.project,
            "code": "TPL-SYS-001",
            "title": "Sistem Gereksinimleri",
            "status": TechnicalDocument.STATUS_IN_REVIEW,
            "revision": "B",
        }
        values.update(overrides)
        document = TechnicalDocument.objects.create(**values)
        document.panels.add(self.panel)
        return document

    def test_admin_creates_document_with_multiple_panel_relation_and_history(self):
        second_panel = ProjectPanel.objects.create(project=self.project, name="Uçuş Kontrol")
        self.client.force_login(self.admin)

        response = self.client.post(
            reverse("technical-document-list"),
            data={
                "project": self.project.id,
                "panels": [self.panel.id, second_panel.id],
                "code": "tpl-icd-002",
                "title": "Arayüz Kontrol Dokümanı",
                "revision": "A",
                "status": "draft",
                "priority": "high",
                "classification": "internal",
            },
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 201)
        document = TechnicalDocument.objects.get(pk=response.json()["id"])
        self.assertEqual(document.code, "TPL-ICD-002")
        self.assertEqual(document.panels.count(), 2)
        self.assertEqual(document.status_history.count(), 1)
        self.assertEqual(document.created_by, self.admin)

    def test_document_rejects_panel_from_another_project(self):
        self.client.force_login(self.admin)
        response = self.client.post(
            reverse("technical-document-list"),
            data={
                "project": self.project.id,
                "panels": [self.other_panel.id],
                "code": "TPL-BAD-001",
                "title": "Geçersiz İlişki",
            },
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("panels", response.json())

    def test_documents_can_share_a_cover_page(self):
        self.client.force_login(self.admin)
        payload = {
            "project": self.project.id,
            "code": "TPL-COV-001",
            "title": "Kapaklı Doküman",
            "cover_page": {"number": "KP-100", "issue": "02"},
        }
        first_response = self.client.post(
            reverse("technical-document-list"),
            data=payload,
            content_type="application/json",
        )
        payload["code"] = "TPL-COV-002"
        second_response = self.client.post(
            reverse("technical-document-list"),
            data=payload,
            content_type="application/json",
        )

        self.assertEqual(first_response.status_code, 201)
        self.assertEqual(second_response.status_code, 201)
        self.assertEqual(CoverPage.objects.count(), 1)
        cover_page = CoverPage.objects.get()
        self.assertEqual(cover_page.technical_documents.count(), 2)
        self.assertEqual(first_response.json()["cover_page"]["number"], "KP-100")
        self.assertEqual(first_response.json()["cover_page"]["issue"], "02")

    def test_regular_user_can_read_but_cannot_create_document(self):
        document = self.create_document()
        self.client.force_login(self.user)

        list_response = self.client.get(
            reverse("technical-document-list"),
            {"project": self.project.id, "status": "in_review"},
        )
        create_response = self.client.post(
            reverse("technical-document-list"),
            data={"project": self.project.id, "code": "TPL-X", "title": "Yetkisiz"},
            content_type="application/json",
        )

        self.assertEqual(list_response.status_code, 200)
        self.assertEqual(list_response.json()[0]["id"], document.id)
        self.assertEqual(create_response.status_code, 403)

    def test_status_update_creates_audit_history(self):
        document = self.create_document()
        self.client.force_login(self.admin)

        response = self.client.patch(
            reverse(
                "technical-document-detail",
                kwargs={"technical_document_id": document.id},
            ),
            data={"status": "approved", "status_note": "Teknik kurul onayı tamamlandı."},
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        history = TechnicalDocumentStatusHistory.objects.get(document=document)
        self.assertEqual(history.from_status, "in_review")
        self.assertEqual(history.to_status, "approved")
        self.assertEqual(history.changed_by, self.admin)

    def test_status_transition_requires_a_non_blank_audit_note(self):
        document = self.create_document()
        self.client.force_login(self.admin)

        response = self.client.patch(
            reverse(
                "technical-document-detail",
                kwargs={"technical_document_id": document.id},
            ),
            data={"status": "approved", "status_note": "   "},
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("status_note", response.json())
        document.refresh_from_db()
        self.assertEqual(document.status, TechnicalDocument.STATUS_IN_REVIEW)
        self.assertFalse(document.status_history.exists())

    def test_project_change_requires_cover_page_reassignment_or_clear(self):
        cover_page = CoverPage.objects.create(
            project=self.project,
            number="KP-OLD",
            issue="01",
        )
        document = self.create_document(cover_page=cover_page)
        self.client.force_login(self.admin)
        detail_url = reverse(
            "technical-document-detail",
            kwargs={"technical_document_id": document.id},
        )

        rejected = self.client.patch(
            detail_url,
            data={
                "project": self.other_project.id,
                "panels": [self.other_panel.id],
            },
            content_type="application/json",
        )

        self.assertEqual(rejected.status_code, 400)
        self.assertIn("cover_page", rejected.json())
        document.refresh_from_db()
        self.assertEqual(document.project, self.project)
        self.assertEqual(document.cover_page, cover_page)

        accepted = self.client.patch(
            detail_url,
            data={
                "project": self.other_project.id,
                "panels": [self.other_panel.id],
                "cover_page": None,
            },
            content_type="application/json",
        )
        self.assertEqual(accepted.status_code, 200)
        document.refresh_from_db()
        self.assertEqual(document.project, self.other_project)
        self.assertIsNone(document.cover_page)

    def test_notification_sends_to_panel_responsibles_and_records_audit(self):
        document = self.create_document()
        self.client.force_login(self.admin)

        response = self.client.post(
            reverse(
                "technical-document-notify",
                kwargs={"technical_document_id": document.id},
            ),
            data={"message": "Dokümanın incelemesi için bilginize."},
            content_type="application/json",
            HTTP_IDEMPOTENCY_KEY="notification-api-0001",
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(len(mail.outbox), 1)
        self.assertEqual(mail.outbox[0].bcc, ["ada@example.com"])
        notification = TechnicalDocumentNotification.objects.get(document=document)
        self.assertEqual(notification.status, "sent")
        self.assertEqual(notification.idempotency_key, "notification-api-0001")
        self.assertEqual(notification.recipient_count, 1)
        document.refresh_from_db()
        self.assertEqual(document.last_notification_recipient_count, 1)

        retry_response = self.client.post(
            reverse(
                "technical-document-notify",
                kwargs={"technical_document_id": document.id},
            ),
            data={"message": "Dokümanın incelemesi için bilginize."},
            content_type="application/json",
            HTTP_IDEMPOTENCY_KEY="notification-api-0001",
        )
        self.assertEqual(retry_response.status_code, 200)
        self.assertIn("daha önce gönderildi", retry_response.json()["message"])
        self.assertEqual(len(mail.outbox), 1)
        self.assertEqual(TechnicalDocumentNotification.objects.filter(document=document).count(), 1)

    def test_notification_requires_a_valid_idempotency_key(self):
        document = self.create_document()
        self.client.force_login(self.admin)

        missing = self.client.post(
            reverse(
                "technical-document-notify",
                kwargs={"technical_document_id": document.id},
            ),
            data={"message": "Eksik anahtar."},
            content_type="application/json",
        )
        invalid = self.client.post(
            reverse(
                "technical-document-notify",
                kwargs={"technical_document_id": document.id},
            ),
            data={"message": "Geçersiz anahtar."},
            content_type="application/json",
            HTTP_IDEMPOTENCY_KEY="short",
        )

        self.assertEqual(missing.status_code, 400)
        self.assertEqual(invalid.status_code, 400)
        self.assertEqual(len(mail.outbox), 0)
        self.assertFalse(TechnicalDocumentNotification.objects.filter(document=document).exists())

    @override_settings(TECHNICAL_NOTIFICATION_PENDING_TIMEOUT=60)
    def test_stale_pending_notification_returns_unknown_without_resending(self):
        document = self.create_document()
        notification = TechnicalDocumentNotification.objects.create(
            document=document,
            subject="Konu",
            message="İçerik",
            recipients=["ada@example.com"],
            recipient_count=1,
            status=TechnicalDocumentNotification.STATUS_PENDING,
            sent_by=self.admin,
            idempotency_key="notification-api-stale",
        )
        TechnicalDocumentNotification.objects.filter(pk=notification.pk).update(
            created_at=timezone.now() - timedelta(minutes=2)
        )
        self.client.force_login(self.admin)

        response = self.client.post(
            reverse(
                "technical-document-notify",
                kwargs={"technical_document_id": document.id},
            ),
            data={"subject": "Konu", "message": "İçerik"},
            content_type="application/json",
            HTTP_IDEMPOTENCY_KEY="notification-api-stale",
        )

        self.assertEqual(response.status_code, 409)
        self.assertEqual(response.json()["notification"]["status"], "unknown")
        self.assertEqual(len(mail.outbox), 0)

    def test_regular_user_cannot_send_technical_document_notification(self):
        document = self.create_document()
        self.client.force_login(self.user)

        response = self.client.post(
            reverse(
                "technical-document-notify",
                kwargs={"technical_document_id": document.id},
            ),
            data={"message": "Yetkisiz bildirim."},
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 403)
        self.assertEqual(len(mail.outbox), 0)
        self.assertFalse(TechnicalDocumentNotification.objects.filter(document=document).exists())

    def test_notification_audit_details_are_visible_only_to_staff(self):
        document = self.create_document()
        TechnicalDocumentNotification.objects.create(
            document=document,
            subject="Gizli konu",
            message="Gizli bildirim gövdesi",
            recipients=["ada@example.com"],
            recipient_count=1,
            status=TechnicalDocumentNotification.STATUS_FAILED,
            error_message="SMTP iç hata ayrıntısı",
            sent_by=self.admin,
        )
        detail_url = reverse(
            "technical-document-detail",
            kwargs={"technical_document_id": document.id},
        )

        self.client.force_login(self.user)
        reader_payload = self.client.get(detail_url).json()
        reader_audit = reader_payload["notifications"][0]
        self.assertEqual(reader_payload["notification_recipients"], [])
        for sensitive_field in ("subject", "message", "recipients", "error_message"):
            self.assertNotIn(sensitive_field, reader_audit)

        self.client.force_login(self.admin)
        staff_payload = self.client.get(detail_url).json()
        staff_audit = staff_payload["notifications"][0]
        self.assertEqual(staff_audit["message"], "Gizli bildirim gövdesi")
        self.assertEqual(staff_audit["recipients"], ["ada@example.com"])
        self.assertEqual(staff_audit["error_message"], "SMTP iç hata ayrıntısı")
        self.assertEqual(
            staff_payload["notification_recipients"][0]["email"],
            "ada@example.com",
        )
