import json
import tempfile
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import SimpleTestCase, override_settings
from docx import Document
from docx.oxml.ns import qn
from docxtpl import DocxTemplate
from pypdf import PdfWriter
from rest_framework.test import APITestCase

from .form_processes.catalog import (
    FORM_TEMPLATES,
    FormTemplateValidationError,
    form_process_catalog,
    validate_form_data,
)
from .models import FormProcessRecord


def valid_pdf_bytes() -> bytes:
    output = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.write(output)
    return output.getvalue()


class FormProcessCatalogTests(SimpleTestCase):
    def test_catalog_contains_every_retained_form_template(self):
        catalog = form_process_catalog()

        self.assertEqual(len(catalog), 6)
        self.assertEqual(len(FORM_TEMPLATES), 36)
        self.assertEqual(len({template.code for template in FORM_TEMPLATES}), 36)
        self.assertEqual(
            sum(len(process["templates"]) for process in catalog),
            len(FORM_TEMPLATES),
        )
        for template in FORM_TEMPLATES:
            with self.subTest(template=template.code):
                self.assertTrue(template.document_path.is_file())
                DocxTemplate(template.document_path)
                self.assertTrue(template.fields)

    def test_flight_permit_forms_are_owned_by_the_engineering_form_catalog(self):
        catalog = form_process_catalog()
        flight_permits = next(process for process in catalog if process["code"] == "flight-permits")
        expected_codes = {
            "fm_dsg_0327",
            "pr_dsg_20_034E",
            "pr_qua_20_104E",
            "fm_qua_0579",
            "fm_qua_0580",
            "fm_qua_0581",
        }

        self.assertEqual(flight_permits["name"], "Uçuş İzinleri")
        self.assertEqual(
            {template["code"] for template in flight_permits["templates"]},
            expected_codes,
        )
        self.assertFalse(
            any(
                template["code"] in expected_codes
                for process in catalog
                if process["code"] != "flight-permits"
                for template in process["templates"]
            )
        )

    def test_catalog_does_not_request_signatures_or_signature_dates(self):
        banned_terms = ("imza", "signature", "signer")
        schema_terms = []
        for template in FORM_TEMPLATES:
            schema_terms.extend((template.title, template.description))
            for field in template.fields:
                schema_terms.extend((field.key, field.label))
                for column in field.columns:
                    schema_terms.extend((column.key, column.label))
            retained_template = DocxTemplate(template.document_path)
            schema_terms.extend(retained_template.get_undeclared_template_variables())

        for term in schema_terms:
            with self.subTest(term=term):
                normalized = term.casefold()
                self.assertFalse(any(banned in normalized for banned in banned_terms))

    def test_removed_signature_fields_are_rejected(self):
        removed_fields_by_template = {
            "fm_dsg_0008e": {
                "employee_signature": "Eski imza girdisi",
                "employee_signature_date": "2026-08-17",
            },
            "fm_dsg_0308e": {"signature_date": "2026-08-17"},
            "fm_dsg_0327": {
                "product_head_date": "2026-08-17",
                "srb_manager_date": "2026-08-17",
            },
        }

        for template_code, data in removed_fields_by_template.items():
            with self.subTest(template=template_code):
                with self.assertRaises(FormTemplateValidationError) as raised:
                    validate_form_data(template_code, data, require_required=False)
                self.assertIn("data", raised.exception.errors)


class FormProcessApiTests(APITestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="form-process-user",
            password="test-password",
            is_active=True,
            first_name="Mehmet",
            last_name="Kaya",
        )
        self.client.force_authenticate(self.user)

    def panel_payload(self, **overrides):
        payload = {
            "template_code": "fm_dsg_0200t",
            "record_number": "panel-2026-001",
            "title": "Uçuş Kontrol Panel Uyum Beyanı",
            "status": "in_review",
            "data": {
                "panel_name": "Uçuş Kontrol Paneli",
                "project_name": "İHA-X",
                "related_documents": "Sistem Sertifikasyon Planı",
                "compliance_documents": "Uyum Matrisi Rev. 2",
                "certification_actions": "Aksiyon bulunmuyor.",
                "declaration": "Panel kapsamındaki gereksinimler karşılanmıştır.",
                "configuration_basis": "Konfigürasyon C-14",
                "panel_coordinator": "Ayşe Yılmaz",
                "panel_members": "Mehmet Kaya",
                "declaration_date": "2026-08-13",
            },
            "notes": "Kurul incelemesine hazır.",
        }
        payload.update(overrides)
        return payload

    def fcc_payload(self, **overrides):
        payload = {
            "template_code": "fm_dsg_0327",
            "record_number": "fcc-2026-001",
            "title": "ANKA Uçuş Uygunluk Belgesi",
            "status": "approved",
            "data": {
                "project_name": "ANKA III",
                "aircraft_type": "İHA",
                "aircraft_model": "ANKA-3 Block A",
                "aircraft_serial_number": "T-001",
                "flight_clearance_reference": "TFCC-01.00",
                "clearance_type": "initial",
                "request_reason": "İlk uçuş ve takip eden geliştirme uçuşları.",
                "request_reason_continued": "Test sahası operasyonları ile sınırlıdır.",
                "valid_from": "2026-08-17",
                "valid_until": "2026-09-30",
                "aircraft_limitation_document_reference": "ALD-ANKA3-001 Rev A",
                "product_head_name": "Ayşe Yılmaz",
                "srb_manager_name": "Mehmet Kaya",
                "statement_of_conformity_reference": "SOC-ANKA3-001",
                "flight_conditions_application_reference": "FCA-ANKA3-001",
                "aircraft_configuration_reference": "ASR-ANKA3-014",
                "airworthiness_safety_summary_documents": "AWSS-001 Rev B\nFSES-004 Rev A",
                "itinerary_airspace_restrictions_reference": "ROUTE-RPT-003",
                "flight_crew_qualification_reference": "FTOM Rev 12",
                "passenger_carrying_restrictions_reference": "OPS-LIM-008",
                "afm_reference": "AFM-ANKA3-DRAFT",
                "flight_test_procedure_reference": "FTP-ANKA3-021",
                "continuing_airworthiness_reference": "ICA-ANKA3-002",
                "clearance_change_criteria_reference": "FCC-CRIT-001",
                "other_appendix": "Silah sistemi emniyet değerlendirmesi WSSA-07.",
                "issue_records": [
                    {
                        "issue": "TFCC-01.00",
                        "date": "2026-08-17",
                        "prepared_by": "Selin Demir",
                        "description": "İlk yayın",
                    }
                ],
            },
            "notes": "SRB kararı ile yayımlandı.",
        }
        payload.update(overrides)
        return payload

    def flight_permit_payload(self, **overrides):
        payload = {
            "template_code": "fm_qua_0579",
            "record_number": "ui-2026-0042",
            "title": "SN-104 Özel Uçuş İzni",
            "status": "approved",
            "data": {
                "applicant": "TUSAŞ",
                "aircraft_owner": "UAV Center",
                "aircraft_model": "Test Platformu",
                "serial_number": "SN-104",
                "purpose_of_flight": ["option_1", "option_6"],
                "purpose_scope": "Geliştirme ve müşteri kabul uçuşu",
                "flight_duration": "3",
                "valid_from": "2026-08-20",
                "valid_until": "2026-10-20",
                "is_recommendation": "no",
            },
            "notes": "Gündüz VFR operasyonları",
        }
        payload.update(overrides)
        return payload

    def test_catalog_requires_active_authentication(self):
        response = self.client.get("/api/form-processes/templates/")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(len(response.data), 6)
        self.assertEqual(
            sum(len(process["templates"]) for process in response.data),
            36,
        )

        self.client.force_authenticate(user=None)
        unauthorized = self.client.get("/api/form-processes/templates/")
        self.assertIn(unauthorized.status_code, {401, 403})

    def test_creates_flight_permit_as_a_shared_engineering_form_record(self):
        response = self.client.post(
            "/api/form-processes/",
            self.flight_permit_payload(),
            format="json",
        )

        self.assertEqual(response.status_code, 201)
        self.assertEqual(response.data["process_code"], "flight-permits")
        self.assertEqual(response.data["form_number"], "FM.QUA.0579")
        self.assertEqual(response.data["data"]["purpose_of_flight"], ["option_1", "option_6"])
        self.assertEqual(FormProcessRecord.objects.get().template_code, "fm_qua_0579")

    def test_rejects_invalid_flight_permit_period_duration_and_purpose(self):
        invalid = self.client.post(
            "/api/form-processes/",
            self.flight_permit_payload(
                data={
                    **self.flight_permit_payload()["data"],
                    "valid_from": "2026-10-20",
                    "valid_until": "2026-08-20",
                    "flight_duration": "0",
                    "purpose_of_flight": ["unknown"],
                }
            ),
            format="json",
        )

        self.assertEqual(invalid.status_code, 400)
        self.assertIn("valid_until", invalid.data)
        self.assertIn("flight_duration", invalid.data)
        self.assertIn("purpose_of_flight", invalid.data)
        self.assertFalse(FormProcessRecord.objects.exists())

    def test_uploads_opens_and_deletes_form_attachment(self):
        with tempfile.TemporaryDirectory(prefix="uav-form-attachment-") as media_root:
            with override_settings(MEDIA_ROOT=media_root):
                payload = self.flight_permit_payload()
                upload = SimpleUploadedFile(
                    "ucus-izni.pdf",
                    valid_pdf_bytes(),
                    content_type="application/pdf",
                )
                response = self.client.post(
                    "/api/form-processes/",
                    {
                        **payload,
                        "data": json.dumps(payload["data"]),
                        "attachment": upload,
                    },
                    format="multipart",
                )

                self.assertEqual(response.status_code, 201)
                self.assertEqual(response.data["attachment_name"], "ucus-izni.pdf")
                self.assertTrue(response.data["attachment_url"].endswith("/attachment/"))
                record = FormProcessRecord.objects.get(pk=response.data["id"])
                stored_path = Path(record.attachment.path)
                self.assertTrue(stored_path.is_file())

                opened = self.client.get(response.data["attachment_url"])
                self.assertEqual(opened.status_code, 200)
                self.assertEqual(opened["Content-Type"], "application/pdf")
                self.assertEqual(opened["X-Content-Type-Options"], "nosniff")
                opened.close()

                with self.captureOnCommitCallbacks(execute=True):
                    deleted = self.client.delete(f"/api/form-processes/{record.pk}/")
                self.assertEqual(deleted.status_code, 204)
                self.assertFalse(stored_path.exists())

    def test_replaces_and_removes_form_attachment_with_storage_cleanup(self):
        with tempfile.TemporaryDirectory(prefix="uav-form-attachment-") as media_root:
            with override_settings(MEDIA_ROOT=media_root):
                payload = self.flight_permit_payload()
                created = self.client.post(
                    "/api/form-processes/",
                    {
                        **payload,
                        "data": json.dumps(payload["data"]),
                        "attachment": SimpleUploadedFile("ilk.pdf", valid_pdf_bytes()),
                    },
                    format="multipart",
                )
                self.assertEqual(created.status_code, 201)
                record = FormProcessRecord.objects.get(pk=created.data["id"])
                first_path = Path(record.attachment.path)

                with self.captureOnCommitCallbacks(execute=True):
                    replaced = self.client.patch(
                        f"/api/form-processes/{record.pk}/",
                        {"attachment": SimpleUploadedFile("yeni.pdf", valid_pdf_bytes())},
                        format="multipart",
                    )
                self.assertEqual(replaced.status_code, 200)
                record.refresh_from_db()
                replacement_path = Path(record.attachment.path)
                self.assertFalse(first_path.exists())
                self.assertTrue(replacement_path.is_file())

                with self.captureOnCommitCallbacks(execute=True):
                    removed = self.client.patch(
                        f"/api/form-processes/{record.pk}/",
                        {"remove_attachment": "true"},
                        format="multipart",
                    )
                self.assertEqual(removed.status_code, 200)
                self.assertEqual(removed.data["attachment_url"], "")
                self.assertFalse(replacement_path.exists())

    def test_create_list_update_and_delete_shared_record(self):
        create_response = self.client.post(
            "/api/form-processes/",
            self.panel_payload(),
            format="json",
        )

        self.assertEqual(create_response.status_code, 201)
        self.assertEqual(create_response.data["process_code"], "others")
        self.assertEqual(create_response.data["record_number"], "PANEL-2026-001")
        self.assertEqual(create_response.data["process_name"], "Others")
        self.assertEqual(create_response.data["form_number"], "FM.DSG.0200T")
        self.assertEqual(
            create_response.data["created_by_name"],
            "Mehmet Kaya (form-process-user)",
        )
        self.assertEqual(
            create_response.data["updated_by_name"],
            "Mehmet Kaya (form-process-user)",
        )
        self.assertEqual(FormProcessRecord.objects.get().created_by, self.user)

        update_response = self.client.patch(
            f"/api/form-processes/{create_response.data['id']}/",
            {"status": "approved", "notes": "Onaylandı."},
            format="json",
        )
        self.assertEqual(update_response.status_code, 200)
        self.assertEqual(update_response.data["status_display"], "Onaylandı")

        list_response = self.client.get("/api/form-processes/?process=others")
        self.assertEqual(list_response.status_code, 200)
        self.assertEqual(len(list_response.data), 1)

        delete_response = self.client.delete(f"/api/form-processes/{create_response.data['id']}/")
        self.assertEqual(delete_response.status_code, 204)
        self.assertFalse(FormProcessRecord.objects.exists())

    def test_shared_record_is_visible_to_active_users_but_not_inactive_users(self):
        create_response = self.client.post(
            "/api/form-processes/",
            self.panel_payload(),
            format="json",
        )
        record_url = f"/api/form-processes/{create_response.data['id']}/"
        other_active_user = get_user_model().objects.create_user(
            username="other-form-process-user",
            password="test-password",
            is_active=True,
        )
        self.client.force_authenticate(other_active_user)

        shared_response = self.client.get(record_url)

        self.assertEqual(shared_response.status_code, 200)
        self.assertEqual(shared_response.data["record_number"], "PANEL-2026-001")

        inactive_user = get_user_model().objects.create_user(
            username="inactive-form-process-user",
            password="test-password",
            is_active=False,
        )
        self.client.force_authenticate(inactive_user)

        inactive_response = self.client.get(record_url)

        self.assertEqual(inactive_response.status_code, 403)

    def test_rejects_missing_unknown_and_invalid_template_fields(self):
        missing_required = self.client.post(
            "/api/form-processes/",
            self.panel_payload(data={}, status="approved"),
            format="json",
        )
        self.assertEqual(missing_required.status_code, 400)
        self.assertIn("panel_name", missing_required.data)

        unknown_field = self.client.post(
            "/api/form-processes/",
            self.panel_payload(
                record_number="PANEL-2026-002",
                data={**self.panel_payload()["data"], "unexpected": "secret"},
            ),
            format="json",
        )
        self.assertEqual(unknown_field.status_code, 400)
        self.assertIn("data", unknown_field.data)

        invalid_date = self.client.post(
            "/api/form-processes/",
            self.panel_payload(
                record_number="PANEL-2026-003",
                data={**self.panel_payload()["data"], "declaration_date": "13/08/2026"},
            ),
            format="json",
        )
        self.assertEqual(invalid_date.status_code, 400)
        self.assertIn("declaration_date", invalid_date.data)
        self.assertFalse(FormProcessRecord.objects.exists())

    def test_draft_allows_missing_required_fields_but_keeps_shape_validation(self):
        draft = self.client.post(
            "/api/form-processes/",
            self.panel_payload(data={}, status="draft"),
            format="json",
        )

        self.assertEqual(draft.status_code, 201)
        self.assertEqual(draft.data["status"], "draft")
        self.assertEqual(draft.data["data"]["panel_name"], "")

        invalid_date = self.client.post(
            "/api/form-processes/",
            self.panel_payload(
                record_number="PANEL-2026-002",
                status="draft",
                data={"declaration_date": "13/08/2026"},
            ),
            format="json",
        )
        self.assertEqual(invalid_date.status_code, 400)
        self.assertIn("declaration_date", invalid_date.data)

        unknown_field = self.client.post(
            "/api/form-processes/",
            self.panel_payload(
                record_number="PANEL-2026-003",
                status="draft",
                data={"unexpected": "value"},
            ),
            format="json",
        )
        self.assertEqual(unknown_field.status_code, 400)
        self.assertIn("data", unknown_field.data)

    def test_fcc_validates_structured_issue_records(self):
        created = self.client.post(
            "/api/form-processes/",
            self.fcc_payload(),
            format="json",
        )
        self.assertEqual(created.status_code, 201)
        self.assertEqual(created.data["data"]["issue_records"][0]["issue"], "TFCC-01.00")

        invalid_data = {
            **self.fcc_payload()["data"],
            "issue_records": [
                {
                    "issue": "TFCC-01.01",
                    "date": "17/08/2026",
                    "prepared_by": "Selin Demir",
                    "description": "Revizyon",
                }
            ],
        }
        invalid = self.client.post(
            "/api/form-processes/",
            self.fcc_payload(record_number="FCC-2026-002", data=invalid_data),
            format="json",
        )
        self.assertEqual(invalid.status_code, 400)
        self.assertIn("issue_records", invalid.data)

        invalid_period = self.client.post(
            "/api/form-processes/",
            self.fcc_payload(
                record_number="FCC-2026-003",
                data={
                    **self.fcc_payload()["data"],
                    "valid_from": "2026-09-30",
                    "valid_until": "2026-08-17",
                },
            ),
            format="json",
        )
        self.assertEqual(invalid_period.status_code, 400)
        self.assertIn("valid_until", invalid_period.data)

    def test_approval_requires_complete_data_and_template_is_immutable(self):
        draft = self.client.post(
            "/api/form-processes/",
            self.panel_payload(data={}, status="draft"),
            format="json",
        )
        record_url = f"/api/form-processes/{draft.data['id']}/"

        incomplete_approval = self.client.patch(record_url, {"status": "approved"}, format="json")
        self.assertEqual(incomplete_approval.status_code, 400)
        self.assertIn("panel_name", incomplete_approval.data)

        template_change = self.client.patch(
            record_url,
            {"template_code": "fm_dsg_0328"},
            format="json",
        )
        self.assertEqual(template_change.status_code, 400)
        self.assertIn("template_code", template_change.data)

    def test_approved_record_can_be_archived_and_reopened_as_draft(self):
        created = self.client.post("/api/form-processes/", self.panel_payload(), format="json")
        record_url = f"/api/form-processes/{created.data['id']}/"

        archived = self.client.patch(record_url, {"status": "archived"}, format="json")
        self.assertEqual(archived.status_code, 200)
        self.assertEqual(archived.data["status"], "archived")

        reopened = self.client.patch(record_url, {"status": "draft"}, format="json")
        self.assertEqual(reopened.status_code, 200)
        self.assertEqual(reopened.data["status"], "draft")

    def test_removed_fcc_signature_dates_are_rejected(self):
        response = self.client.post(
            "/api/form-processes/",
            self.fcc_payload(
                data={
                    **self.fcc_payload()["data"],
                    "product_head_date": "2026-08-17",
                    "srb_manager_date": "2026-08-17",
                }
            ),
            format="json",
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("data", response.data)

    def test_record_number_is_unique_within_process(self):
        first = self.client.post("/api/form-processes/", self.panel_payload(), format="json")
        duplicate = self.client.post(
            "/api/form-processes/",
            self.panel_payload(title="İkinci kayıt"),
            format="json",
        )

        self.assertEqual(first.status_code, 201)
        self.assertEqual(duplicate.status_code, 400)
        self.assertIn("record_number", duplicate.data)

    def test_generates_word_document_from_retained_template_and_validated_data(self):
        create_response = self.client.post(
            "/api/form-processes/",
            self.panel_payload(),
            format="json",
        )

        response = self.client.get(
            f"/api/form-processes/{create_response.data['id']}/generated-document/"
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertEqual(response["X-Content-Type-Options"], "nosniff")
        generated_bytes = b"".join(response.streaming_content)
        generated = Document(BytesIO(generated_bytes))
        text_parts = [paragraph.text for paragraph in generated.paragraphs]
        for table in generated.tables:
            text_parts.extend(cell.text for row in table.rows for cell in row.cells)
        generated_text = "\n".join(text_parts)
        self.assertIn("PANEL UYUM BEYANI", generated_text.upper())
        self.assertIn("SÜREÇ KAYIT BİLGİLERİ", generated_text)
        self.assertIn("Uçuş Kontrol Paneli", generated_text)
        self.assertIn("PANEL-2026-001", generated_text)
        self.assertNotIn("{{", generated_text)
        with ZipFile(BytesIO(generated_bytes)) as package:
            self.assertIn("word/document.xml", package.namelist())

    def test_generates_fcc_document_from_authoritative_template(self):
        created = self.client.post(
            "/api/form-processes/",
            self.fcc_payload(),
            format="json",
        )
        self.assertEqual(created.status_code, 201)

        response = self.client.get(f"/api/form-processes/{created.data['id']}/generated-document/")
        self.assertEqual(response.status_code, 200)
        generated_bytes = b"".join(response.streaming_content)
        generated = Document(BytesIO(generated_bytes))
        text_parts = [node.text or "" for node in generated._element.iter(qn("w:t"))]
        for section in generated.sections:
            for header in (section.header, section.first_page_header):
                text_parts.extend(node.text or "" for node in header._element.iter(qn("w:t")))
        generated_text = "".join(text_parts)

        self.assertIn("FLIGHT CLEARANCE CERTIFICATE", generated_text)
        self.assertIn("ANKA III", generated_text)
        self.assertIn("T-001", generated_text)
        self.assertIn("SOC-ANKA3-001", generated_text)
        self.assertIn("17.08.2026", generated_text)
        self.assertIn("TFCC-01.00", generated_text)
        self.assertIn("☒", generated_text)
        self.assertNotIn("{{", generated_text)
        self.assertNotIn("e.g. First Flight", generated_text)
        self.assertNotIn("This Section should contain", generated_text)
        signature_table = generated.tables[3].rows[1].cells[0].tables[0]
        self.assertEqual(signature_table.rows[1].cells[0].text.strip(), "Date:")
        self.assertEqual(signature_table.rows[1].cells[1].text.strip(), "Date:")
