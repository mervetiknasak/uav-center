"""Build the versioned DOCX template used by docxtpl for flight permits."""

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = "0F172A"
MUTED = "64748B"
ACCENT = "0F766E"
LIGHT_FILL = "F1F5F9"
BORDER = "CBD5E1"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
LABEL_WIDTH_DXA = 2880
VALUE_WIDTH_DXA = TABLE_WIDTH_DXA - LABEL_WIDTH_DXA


def set_run_font(run, *, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, *, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, column_widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    table_w = table_pr.find(qn("w:tblW"))
    table_w.set(qn("w:w"), str(sum(column_widths)))
    table_w.set(qn("w:type"), "dxa")

    table_indent = table_pr.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_pr.append(table_indent)
    table_indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    table_indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in column_widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, column_widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Sayfa ")
    set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Title", 24, INK, 0, 4),
        ("Subtitle", 12, MUTED, 0, 14),
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.font.underline = False
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style_ppr = style._element.get_or_add_pPr()
        paragraph_border = style_ppr.find(qn("w:pBdr"))
        if paragraph_border is not None:
            style_ppr.remove(paragraph_border)


def add_detail_row(table, label, placeholder):
    cells = table.add_row().cells
    set_cell_fill(cells[0], LIGHT_FILL)
    set_cell_fill(cells[1], WHITE)
    label_paragraph = cells[0].paragraphs[0]
    label_paragraph.paragraph_format.space_after = Pt(0)
    label_run = label_paragraph.add_run(label)
    set_run_font(label_run, size=9.5, bold=True, color=MUTED)
    value_paragraph = cells[1].paragraphs[0]
    value_paragraph.paragraph_format.space_after = Pt(0)
    value_run = value_paragraph.add_run(placeholder)
    set_run_font(value_run, size=10.5, bold=True, color=INK)


def build_template(output_path):
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    configure_styles(document)
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.paragraph_format.space_after = Pt(0)
    header_run = header_paragraph.add_run("UAV CENTER  |  UÇUŞ OPERASYONLARI")
    set_run_font(header_run, size=8.5, bold=True, color=ACCENT)

    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(footer_table, [7200, 2160])
    set_repeat_table_header(footer_table.rows[0])
    footer_table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)
    footer_note = (
        footer_table.rows[0]
        .cells[0]
        .paragraphs[0]
        .add_run("Bu belge UAV Center uçuş izni kaydından elektronik olarak oluşturulmuştur.")
    )
    set_run_font(footer_note, size=8, color=MUTED)
    add_page_number(footer_table.rows[0].cells[1].paragraphs[0])

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(6)
    kicker_run = kicker.add_run("RESMİ İZİN BELGESİ")
    set_run_font(kicker_run, size=9, bold=True, color=ACCENT)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("UÇUŞ İZNİ")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("FLIGHT PERMIT")

    control = document.add_paragraph()
    control.alignment = WD_ALIGN_PARAGRAPH.CENTER
    control.paragraph_format.space_after = Pt(16)
    control_run = control.add_run("Belge No: {{ permit_number }}")
    set_run_font(control_run, size=10, bold=True, color=MUTED)

    lead = document.add_paragraph()
    lead.paragraph_format.space_after = Pt(12)
    lead.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    lead_run = lead.add_run(
        "Bu belge, aşağıda bilgileri bulunan hava aracının belirtilen tarih aralığı ve "
        "uçuş bölgesi kapsamında operasyon yürütme iznini kayıt altına alır."
    )
    set_run_font(lead_run, size=10.5, color=INK)

    details = document.add_table(rows=1, cols=2)
    details.style = "Table Grid"
    header_cells = details.rows[0].cells
    set_cell_fill(header_cells[0], ACCENT)
    set_cell_fill(header_cells[1], ACCENT)
    header_cells[0].merge(header_cells[1])
    header_paragraph = details.rows[0].cells[0].paragraphs[0]
    header_paragraph.paragraph_format.space_after = Pt(0)
    header_run = header_paragraph.add_run("İZİN BİLGİLERİ / PERMIT DETAILS")
    set_run_font(header_run, size=10, bold=True, color=WHITE)
    set_repeat_table_header(details.rows[0])

    add_detail_row(details, "UÇAK NUMARASI", "{{ aircraft_number }}")
    add_detail_row(details, "UÇUŞ İZİN NUMARASI", "{{ permit_number }}")
    add_detail_row(details, "İZİN TÜRÜ", "{{ permit_type }}")
    add_detail_row(details, "İZNİ VEREN KURUM", "{{ issuing_authority }}")
    add_detail_row(details, "UÇUŞ BÖLGESİ / KAPSAM", "{{ flight_region }}")
    add_detail_row(details, "GEÇERLİLİK BAŞLANGICI", "{{ valid_from }}")
    add_detail_row(details, "GEÇERLİLİK BİTİŞİ", "{{ valid_until }}")
    add_detail_row(details, "İZİN DURUMU", "{{ validity_status }}")
    set_table_geometry(details, [LABEL_WIDTH_DXA, VALUE_WIDTH_DXA])

    notes_heading = document.add_paragraph()
    notes_heading.paragraph_format.space_before = Pt(12)
    notes_heading.paragraph_format.space_after = Pt(4)
    notes_heading_run = notes_heading.add_run("AÇIKLAMALAR / NOTES")
    set_run_font(notes_heading_run, size=9.5, bold=True, color=ACCENT)
    notes = document.add_paragraph()
    notes.paragraph_format.space_after = Pt(14)
    notes_run = notes.add_run("{{ notes }}")
    set_run_font(notes_run, size=10, color=INK)

    approval = document.add_paragraph()
    approval.paragraph_format.space_before = Pt(8)
    approval.paragraph_format.space_after = Pt(20)
    approval_run = approval.add_run(
        "Bu izin yalnızca belirtilen kapsam ve geçerlilik tarihleri içinde kullanılabilir. "
        "Uçuş öncesinde güncel operasyonel ve yasal gerekliliklerin kontrolü sorumluluğu kullanıcıya aittir."
    )
    set_run_font(approval_run, size=9, italic=True, color=MUTED)

    signature_table = document.add_table(rows=2, cols=2)
    signature_table.style = "Table Grid"
    for cell in signature_table.rows[0].cells:
        set_cell_fill(cell, LIGHT_FILL)
    set_repeat_table_header(signature_table.rows[0])
    signature_labels = (("DÜZENLEME TARİHİ", "{{ generated_at }}"), ("ONAY / İMZA", "Yetkili İmza"))
    for index, (label, value) in enumerate(signature_labels):
        label_p = signature_table.rows[0].cells[index].paragraphs[0]
        label_p.paragraph_format.space_after = Pt(0)
        set_run_font(label_p.add_run(label), size=9, bold=True, color=MUTED)
        value_p = signature_table.rows[1].cells[index].paragraphs[0]
        value_p.paragraph_format.space_before = Pt(16 if index else 5)
        value_p.paragraph_format.space_after = Pt(5)
        set_run_font(value_p.add_run(value), size=10, bold=not index, color=INK)
    set_table_geometry(signature_table, [4680, 4680])

    document.core_properties.title = "Uçuş İzni / Flight Permit"
    document.core_properties.subject = "Uçak bazlı uçuş izin belgesi"
    document.core_properties.author = "UAV Center"
    document.core_properties.keywords = "uçuş izni, flight permit, UAV"
    document.save(output_path)
    return output_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build_template(args.output)
