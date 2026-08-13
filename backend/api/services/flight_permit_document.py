from io import BytesIO

from django.utils import timezone
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxtpl import DocxTemplate

from ..flight_permits.purposes import flight_purpose_labels
from ..flight_permits.templates import get_flight_permit_template


def _format_date(value):
    return value.strftime("%d.%m.%Y") if value else "-"


def _join_values(*values):
    return " / ".join(str(value).strip() for value in values if str(value or "").strip()) or "-"


def _target_date_and_duration(permit):
    duration = f"{permit.flight_duration} saat" if permit.flight_duration else ""
    return _join_values(_format_date(permit.target_date) if permit.target_date else "", duration)


def _set_summary_table_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:color"), "B7C6D8")
        borders.append(border)
    table._tbl.tblPr.append(borders)


def build_flight_permit_document(permit):
    template_definition = get_flight_permit_template(permit.template_code)
    template = DocxTemplate(template_definition.document_path)
    context = {
        "institution": template_definition.institution,
        "permit_applicant": permit.permit_applicant,
        "permit_number": permit.permit_number,
        "aircraft_nationality": _join_values(
            permit.aircraft_nationality,
            permit.aircraft_id_mark,
        ),
        "aircraft_id_mark": "",
        "aircraft_owner": permit.aircraft_owner or "-",
        "aircraft_manufacturer": _join_values(
            permit.aircraft_manufacturer,
            permit.aircraft_type,
        ),
        "aircraft_type": "",
        "serial_number": permit.serial_number or "-",
        "purpose_of_flight": "  •  ".join(flight_purpose_labels(permit.purpose_of_flight)) or "-",
        "target_date": _target_date_and_duration(permit),
        "flight_duration": "",
        "aircraft_configuration": permit.aircraft_configuration or "-",
        "conditions_restrictions": permit.conditions_restrictions or "-",
        "conditions_substantiations": permit.conditions_substantiations or "-",
        "is_recommendation": permit.is_recommendation,
        "valid_from": _format_date(permit.valid_from),
        "valid_until": _format_date(permit.valid_until),
        "generated_at": _format_date(timezone.localdate()),
        **permit.template_data,
    }
    template.render(context, autoescape=True)
    output = BytesIO()
    template.save(output)
    if template_definition.append_field_summary:
        output.seek(0)
        document = Document(output)
        heading = document.add_paragraph()
        heading.paragraph_format.keep_with_next = True
        heading.add_run(f"{template_definition.institution} — Kuruma Özel Bilgiler").bold = True
        summary = document.add_table(rows=0, cols=2)
        _set_summary_table_borders(summary)
        for field in template_definition.fields:
            cells = summary.add_row().cells
            cells[0].text = field.label
            cells[1].text = permit.template_data.get(field.key) or "-"
        output = BytesIO()
        document.save(output)
    output.seek(0)
    return output
