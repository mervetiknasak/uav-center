from dataclasses import asdict, dataclass
from pathlib import Path


class WordTableParseError(ValueError):
    pass


@dataclass(frozen=True)
class WordTableCell:
    index: int
    table_index: int
    row_index: int
    column_index: int
    text: str

    def as_dict(self):
        return asdict(self)


FIELD_COORDINATES = {
    "project": (0, 1, 2),
    "subject": (0, 2, 2),
    "mom_no": (0, 3, 2),
    "revision": (0, 4, 2),
    "date_time": (0, 5, 2),
    "location": (0, 6, 2),
    "agenda": (0, 8, 0),
    "discussions_decisions": (0, 10, 0)
}


def _extract_mapped_data(cells):
    coordinates = {
        (cell.table_index, cell.row_index, cell.column_index): cell.text
        for cell in cells
    }
    fields = {
        field_name: coordinates.get(field_coordinates, "")
        for field_name, field_coordinates in FIELD_COORDINATES.items()
    }

    action_start = next(
        (
            index
            for index, cell in enumerate(cells)
            if any(
                heading in cell.text.casefold()
                for heading in ("action item list", "aksiyon listesi", "aksiyon maddeleri")
            )
        ),
        None,
    )
    attachments_start = next(
        (
            index
            for index, cell in enumerate(cells)
            if action_start is not None
            and index > action_start
            and any(
                heading in cell.text.casefold()
                for heading in ("attachments", "ekler")
            )
        ),
        None,
    )

    action_items = []
    if action_start is not None:
        item_cells = cells[action_start + 5 : attachments_start]
        for offset in range(0, len(item_cells), 4):
            group = item_cells[offset : offset + 4]
            if len(group) < 4:
                break
            values = [cell.text for cell in group]
            if not any(values):
                continue

            action_items.append(
                {
                    "no": values[0],
                    "action_item": values[1],
                    "responsible": values[2],
                    "due_date": values[3],
                }
            )

    return {
        **fields,
        "action_items": action_items,
        "action_item_list_found": action_start is not None,
        "attachments_found": attachments_start is not None,
    }


def parse_word_table(file_path):
    """Return every visible DOCX table cell in deterministic document order."""
    from docx import Document

    path = Path(file_path)
    if path.suffix.lower() != ".docx":
        raise WordTableParseError("Yalnızca .docx uzantılı Word dosyaları destekleniyor.")

    try:
        document = Document(str(path))
    except Exception as exc:
        raise WordTableParseError(f"Word dosyası açılamadı: {exc}") from exc

    if not document.tables:
        raise WordTableParseError("Word dosyasında tablo bulunamadı.")

    cells = []
    seen_cells = set()
    global_index = 0
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)

                text = "\n".join(
                    line.strip()
                    for line in cell.text.splitlines()
                    if line.strip()
                )
                parsed_cell = WordTableCell(
                    index=global_index,
                    table_index=table_index,
                    row_index=row_index,
                    column_index=column_index,
                    text=text,
                )
                cells.append(parsed_cell)
                print(
                    "[WORD_CELL] "
                    f"index={global_index} table={table_index} "
                    f"row={row_index} column={column_index} text={text!r}",
                    flush=True,
                )
                global_index += 1

    return {
        "table_count": len(document.tables),
        "cell_count": len(cells),
        "cells": [cell.as_dict() for cell in cells],
        "extracted_data": _extract_mapped_data(cells),
    }
