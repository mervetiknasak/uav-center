"""Render a validated record with its retained FM DOCX template."""

from datetime import date
from io import BytesIO

from django.utils import timezone
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docxtpl import DocxTemplate

from ..catalog import get_form_template


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_table_borders(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "5")
        border.set(qn("w:color"), "B7C6D8")
        borders.append(border)
    table._tbl.tblPr.append(borders)


def _display_value(value, field=None) -> str:
    if value is True:
        return "Evet"
    if value is False:
        return "Hayır"
    if isinstance(value, list):
        if not value:
            return "—"
        if getattr(field, "field_type", "") == "multi_select":
            labels = dict(getattr(field, "options", ()))
            return "\n".join(str(labels.get(item, item) or "") for item in value)
        columns = getattr(field, "columns", ())
        return (
            "\n".join(
                " | ".join(
                    f"{column.label}: "
                    f"{_display_date(row.get(column.key)) if column.field_type == 'date' else row.get(column.key) or '—'}"
                    for column in columns
                )
                for row in value
                if isinstance(row, dict)
            )
            or "—"
        )
    if getattr(field, "field_type", "") == "date":
        return _display_date(value) or "—"
    if getattr(field, "field_type", "") == "select":
        return dict(field.options).get(value, value) or "—"
    return str(value or "—")


def _display_date(value) -> str:
    if not value:
        return ""
    try:
        return date.fromisoformat(value).strftime("%d.%m.%Y")
    except (TypeError, ValueError):
        return str(value)


def _template_context(record, definition) -> dict:
    context = {
        "record_number": record.record_number,
        "record_title": record.title,
        "record_status": record.get_status_display(),
        "generated_at": timezone.localdate().strftime("%d.%m.%Y"),
        **record.data,
    }
    if definition.code != "fm_dsg_0327":
        return context

    clearance_type = record.data.get("clearance_type")
    context.update(
        {
            "clearance_initial_mark": "☒" if clearance_type == "initial" else "☐",
            "clearance_renewal_mark": "☒" if clearance_type == "renewal" else "☐",
            "clearance_cancelled_mark": ("☒" if clearance_type == "cancelled_suspended" else "☐"),
            "valid_from_display": _display_date(record.data.get("valid_from")),
            "valid_until_display": _display_date(record.data.get("valid_until")),
        }
    )
    issue_records = []
    for row in record.data.get("issue_records") or []:
        normalized = dict(row)
        normalized["date_display"] = _display_date(row.get("date"))
        issue_records.append(normalized)
    empty_issue = {
        "issue": "",
        "date": "",
        "date_display": "",
        "prepared_by": "",
        "description": "",
    }
    context["issue_records"] = (issue_records + [empty_issue] * 5)[:5]
    return context


def build_form_process_document(record):
    definition = get_form_template(record.template_code)
    template = DocxTemplate(definition.document_path)
    context = _template_context(record, definition)
    template.render(context, autoescape=True)
    rendered = BytesIO()
    template.save(rendered)
    rendered.seek(0)

    document = Document(rendered)
    document.add_page_break()
    heading = document.add_paragraph()
    heading.paragraph_format.keep_with_next = True
    heading.add_run("SÜREÇ KAYIT BİLGİLERİ").bold = True
    metadata = document.add_table(rows=0, cols=2)
    _set_table_borders(metadata)
    metadata_rows = (
        ("Süreç", definition.process_name),
        ("Form", f"{definition.form_number} — {definition.title}"),
        ("Kayıt numarası", record.record_number),
        ("Kayıt başlığı", record.title),
        ("Durum", record.get_status_display()),
        ("Oluşturma tarihi", timezone.localtime(record.created_at).strftime("%d.%m.%Y %H:%M")),
        ("Son güncelleyen", getattr(record.updated_by, "username", "") or "—"),
    )
    for label, value in metadata_rows:
        cells = metadata.add_row().cells
        _set_cell_shading(cells[0], "E8EEF6")
        cells[0].text = label
        cells[1].text = _display_value(value)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    current_group = None
    details: Table | None = None
    for field in definition.fields:
        if field.group != current_group:
            current_group = field.group
            group_heading = document.add_paragraph()
            group_heading.paragraph_format.keep_with_next = True
            group_heading.add_run(current_group).bold = True
            details = document.add_table(rows=0, cols=2)
            _set_table_borders(details)
        if details is None:
            raise RuntimeError("Form alan grubu tablosu oluşturulamadı.")
        cells = details.add_row().cells
        _set_cell_shading(cells[0], "F1F5F9")
        cells[0].text = field.label
        cells[1].text = _display_value(record.data.get(field.key), field)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    if record.notes:
        notes_heading = document.add_paragraph()
        notes_heading.add_run("Notlar").bold = True
        document.add_paragraph(record.notes)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output
