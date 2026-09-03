import logging
import unicodedata
from dataclasses import asdict, dataclass
from pathlib import Path

from ..common.redaction import safe_exception_message
from ..services.document_limits import DocumentPreflightError, validate_office_archive

logger = logging.getLogger(__name__)


class EDKMinutesParseError(ValueError):
    pass


@dataclass(frozen=True)
class EDKMinutesCell:
    index: int
    table_index: int
    row_index: int
    column_index: int
    text: str

    def as_dict(self):
        return asdict(self)


@dataclass(frozen=True)
class EDKMinutesTextBlock:
    text: str


FIELD_COORDINATES = {
    "project": (0, 1, 2),
    "subject": (0, 2, 2),
    "mom_no": (0, 3, 2),
    "revision": (0, 4, 2),
    "date_time": (0, 5, 2),
    "location": (0, 6, 2),
    "agenda": (0, 8, 0),
    "discussions_decisions": (0, 10, 0),
}

DISCUSSIONS_HEADINGS = (
    "discussions and decisions",
    "discussions & decisions",
    "topics discussed",
    "görüşmeler ve kararlar",
    "görüşülen konular",
)
ACTION_ITEM_HEADINGS = (
    "action item list",
    "aksiyon listesi",
    "aksiyon maddeleri",
    "işlem maddeleri",
)
ATTACHMENTS_HEADINGS = ("attachments", "ekler")


def _normalized_text(text):
    return "".join(
        character
        for character in unicodedata.normalize("NFKD", text.casefold().replace("ı", "i"))
        if not unicodedata.combining(character)
    )


def _heading_position(text, headings):
    offset = 0
    for line in text.splitlines(keepends=True) or (text,):
        folded_line = _normalized_text(line)
        for heading in headings:
            folded_heading = _normalized_text(heading)
            heading_start = folded_line.find(folded_heading)
            if heading_start < 0:
                continue

            prefix = folded_line[:heading_start]
            if any(character.isalpha() for character in prefix):
                continue

            return offset + heading_start, folded_heading
        offset += len(line)
    return None


def _find_heading_index(cells, headings, *, start=0):
    return next(
        (
            index
            for index, cell in enumerate(cells[start:], start=start)
            if _heading_position(cell.text, headings) is not None
        ),
        None,
    )


def _text_after_heading(text, headings):
    position = _heading_position(text, headings)
    if position is None:
        return text.strip()

    heading_start, heading = position
    content_start = heading_start + len(heading)
    return text[content_start:].lstrip(" \t\r\n:;-–—").strip()


def _extract_section(cells, *, start_headings, end_index):
    start_index = _find_heading_index(cells, start_headings)
    if start_index is None or end_index is None or start_index >= end_index:
        return None

    values = []
    for index, cell in enumerate(cells[start_index:end_index], start=start_index):
        text = cell.text.strip()
        if index == start_index or _heading_position(text, start_headings) is not None:
            text = _text_after_heading(text, start_headings)
        if text:
            values.append(text)

    return "\n".join(values)


def _extract_mapped_data(cells, content_blocks=None):
    content_blocks = cells if content_blocks is None else content_blocks
    coordinates = {
        (cell.table_index, cell.row_index, cell.column_index): cell.text for cell in cells
    }
    fields = {
        field_name: coordinates.get(field_coordinates, "")
        for field_name, field_coordinates in FIELD_COORDINATES.items()
    }

    action_start = _find_heading_index(cells, ACTION_ITEM_HEADINGS)
    attachments_start = (
        _find_heading_index(cells, ATTACHMENTS_HEADINGS, start=action_start + 1)
        if action_start is not None
        else None
    )

    content_action_start = _find_heading_index(content_blocks, ACTION_ITEM_HEADINGS)
    discussions_decisions = _extract_section(
        content_blocks,
        start_headings=DISCUSSIONS_HEADINGS,
        end_index=content_action_start,
    )
    if discussions_decisions is not None:
        fields["discussions_decisions"] = discussions_decisions

    action_items = []
    if action_start is not None:
        item_cells = cells[action_start + 5 : attachments_start]
        for offset in range(0, len(item_cells), 4):
            group = item_cells[offset : offset + 4]
            if len(group) < 4:
                break
            values = [cell.text for cell in group]
            if not any(values):
                continue

            action_items.append(
                {
                    "no": values[0],
                    "action_item": values[1],
                    "responsible": values[2],
                    "due_date": values[3],
                }
            )

    return {
        **fields,
        "action_items": action_items,
        "action_item_list_found": action_start is not None,
        "attachments_found": attachments_start is not None,
    }


def parse_minutes_document(file_path):
    """Return every visible DOCX table cell in deterministic document order."""
    from docx import Document
    from docx.table import Table

    path = Path(file_path)
    if path.suffix.lower() != ".docx":
        raise EDKMinutesParseError("Yalnızca .docx uzantılı Word dosyaları destekleniyor.")

    try:
        validate_office_archive(path, ".docx")
        document = Document(str(path))
    except DocumentPreflightError as exc:
        raise EDKMinutesParseError(str(exc)) from exc
    except Exception as exc:
        logger.warning(
            "Word document open failed: %s",
            safe_exception_message(exc),
            extra={"event": "word_document_open_failed"},
        )
        raise EDKMinutesParseError("Word dosyası açılamadı.") from exc

    if not document.tables:
        raise EDKMinutesParseError("Word dosyasında tablo bulunamadı.")

    cells = []
    content_blocks = []
    seen_cells = set()
    global_index = 0
    table_index = 0
    for block in document.iter_inner_content():
        if not isinstance(block, Table):
            text = "\n".join(line.strip() for line in block.text.splitlines() if line.strip())
            if text:
                content_blocks.append(EDKMinutesTextBlock(text=text))
            continue

        table = block
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)

                text = "\n".join(line.strip() for line in cell.text.splitlines() if line.strip())
                parsed_cell = EDKMinutesCell(
                    index=global_index,
                    table_index=table_index,
                    row_index=row_index,
                    column_index=column_index,
                    text=text,
                )
                cells.append(parsed_cell)
                if text:
                    content_blocks.append(EDKMinutesTextBlock(text=text))
                global_index += 1
        table_index += 1

    return {
        "table_count": len(document.tables),
        "cell_count": len(cells),
        "cells": [cell.as_dict() for cell in cells],
        "extracted_data": _extract_mapped_data(cells, content_blocks),
    }
